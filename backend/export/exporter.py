"""
export/exporter.py
------------------
Exports meeting summaries and transcripts to various document formats.

Supported formats:
  - txt   : Plain text (always generated as a baseline)
  - docx  : Microsoft Word (.docx) — rich formatting, headers, tables
  - pdf   : PDF via reportlab — portable, printable
  - all   : Generate all three formats

Requirements:
    pip install python-docx reportlab
"""

from __future__ import annotations

import os
import datetime
from pathlib import Path
from typing import Optional
from utils.logger import get_logger

logger = get_logger(__name__)


class Exporter:
    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export(
        self,
        fmt: str,
        transcript: str,
        summary: str,
        speakers: list[str],
        duration: float,
        timestamp: str,
    ) -> str:
        """
        Export to the specified format. Returns path to the primary output file.
        For 'all', returns the DOCX path.
        """
        meta = MeetingMeta(
            timestamp=timestamp,
            duration_seconds=duration,
            speakers=speakers,
        )

        paths = []
        if fmt in ("txt", "all"):
            paths.append(self._export_txt(transcript, summary, meta))
        if fmt in ("docx", "all"):
            paths.append(self._export_docx(transcript, summary, meta))
        if fmt in ("pdf", "all"):
            paths.append(self._export_pdf(transcript, summary, meta))

        if not paths:
            raise ValueError(f"Unknown export format: {fmt}")

        return paths[0]

    # ------------------------------------------------------------------
    # TXT
    # ------------------------------------------------------------------

    def _export_txt(self, transcript: str, summary: str, meta: "MeetingMeta") -> str:
        path = self.output_dir / "meeting_report.txt"
        with open(path, "w", encoding="utf-8") as f:
            f.write(_TXT_HEADER.format(
                date=meta.date_str,
                duration=meta.duration_str,
                speakers=", ".join(meta.speakers) if meta.speakers else "Unknown",
            ))
            f.write("\n\n")
            f.write("=" * 60 + "\n")
            f.write("SUMMARY\n")
            f.write("=" * 60 + "\n\n")
            f.write(summary)
            f.write("\n\n")
            f.write("=" * 60 + "\n")
            f.write("FULL TRANSCRIPT\n")
            f.write("=" * 60 + "\n\n")
            f.write(transcript)
        logger.info(f"TXT exported: {path}")
        return str(path)

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    def _export_docx(self, transcript: str, summary: str, meta: "MeetingMeta") -> str:
        path = self.output_dir / "meeting_report.docx"
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise RuntimeError(
                "python-docx is not installed. Run: pip install python-docx"
            )

        doc = Document()

        # ── Document styles ────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # ── Cover / header ─────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("Meeting Summary")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        doc.add_paragraph()

        meta_table = doc.add_table(rows=3, cols=2)
        meta_table.style = "Table Grid"
        _set_cell(meta_table, 0, 0, "Date", bold=True)
        _set_cell(meta_table, 0, 1, meta.date_str)
        _set_cell(meta_table, 1, 0, "Duration", bold=True)
        _set_cell(meta_table, 1, 1, meta.duration_str)
        _set_cell(meta_table, 2, 0, "Speakers", bold=True)
        _set_cell(meta_table, 2, 1, ", ".join(meta.speakers) if meta.speakers else "Unknown")

        doc.add_paragraph()
        doc.add_paragraph()

        # ── Summary section ────────────────────────────────────────────
        _add_section_heading(doc, "Summary")
        _add_markdown_content(doc, summary)

        doc.add_page_break()

        # ── Transcript section ─────────────────────────────────────────
        _add_section_heading(doc, "Full Transcript")
        for line in transcript.split("\n"):
            if not line.strip():
                continue
            p = doc.add_paragraph()
            # Bold speaker labels like "[00:01] Speaker 1:"
            if ":" in line and line.startswith("["):
                parts = line.split(":", 2)
                if len(parts) >= 2:
                    label = parts[0] + ":"
                    content = parts[1] if len(parts) == 2 else ":".join(parts[1:])
                    p.add_run(label).bold = True
                    p.add_run(content)
                else:
                    p.add_run(line)
            else:
                p.add_run(line)
            p.paragraph_format.space_after = Pt(3)

        doc.save(str(path))
        logger.info(f"DOCX exported: {path}")
        return str(path)

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _export_pdf(self, transcript: str, summary: str, meta: "MeetingMeta") -> str:
        path = self.output_dir / "meeting_report.pdf"
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer,
                Table, TableStyle, PageBreak, HRFlowable,
            )
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
        except ImportError:
            raise RuntimeError(
                "reportlab is not installed. Run: pip install reportlab"
            )

        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=2.5 * cm,
            rightMargin=2.5 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm,
        )

        styles = getSampleStyleSheet()
        brand_blue = colors.HexColor("#1A56DB")

        title_style = ParagraphStyle(
            "MeetTitle",
            parent=styles["Title"],
            fontSize=26,
            textColor=brand_blue,
            spaceAfter=12,
            alignment=TA_CENTER,
        )
        h2_style = ParagraphStyle(
            "MeetH2",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=brand_blue,
            spaceBefore=14,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            "MeetBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=15,
            spaceAfter=4,
        )
        mono_style = ParagraphStyle(
            "MeetMono",
            parent=styles["Code"],
            fontSize=9,
            leading=13,
            spaceAfter=2,
        )

        story = []

        # Title
        story.append(Paragraph("Meeting Summary", title_style))
        story.append(Spacer(1, 0.3 * cm))

        # Meta table
        meta_data = [
            ["Date", meta.date_str],
            ["Duration", meta.duration_str],
            ["Speakers", ", ".join(meta.speakers) if meta.speakers else "Unknown"],
        ]
        t = Table(meta_data, colWidths=[4 * cm, 12 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EBF5FB")),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ROWBACKGROUNDS", (1, 0), (-1, -1), [colors.white, colors.HexColor("#F8F9FA")]),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.grey),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.6 * cm))

        # Summary
        story.append(Paragraph("Summary", h2_style))
        story.append(HRFlowable(width="100%", thickness=1, color=brand_blue))
        story.append(Spacer(1, 0.2 * cm))
        for line in summary.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.15 * cm))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], h2_style))
            elif line.startswith("- "):
                story.append(Paragraph(f"• {line[2:]}", body_style))
            else:
                story.append(Paragraph(line, body_style))

        story.append(PageBreak())

        # Transcript
        story.append(Paragraph("Full Transcript", h2_style))
        story.append(HRFlowable(width="100%", thickness=1, color=brand_blue))
        story.append(Spacer(1, 0.2 * cm))
        for line in transcript.split("\n"):
            if line.strip():
                # Escape HTML special chars for reportlab
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, mono_style))

        doc.build(story)
        logger.info(f"PDF exported: {path}")
        return str(path)


# ------------------------------------------------------------------
# Data classes & helpers
# ------------------------------------------------------------------

class MeetingMeta:
    def __init__(self, timestamp: str, duration_seconds: float, speakers: list[str]):
        self.speakers = speakers
        self.duration_seconds = duration_seconds
        try:
            dt = datetime.datetime.strptime(timestamp, "%Y%m%d_%H%M%S")
            self.date_str = dt.strftime("%B %d, %Y at %I:%M %p")
        except Exception:
            self.date_str = timestamp

    @property
    def duration_str(self) -> str:
        m, s = divmod(int(self.duration_seconds), 60)
        h, m = divmod(m, 60)
        if h:
            return f"{h}h {m}m {s}s"
        return f"{m}m {s}s"


_TXT_HEADER = """\
MEETING SUMMARY
===============
Date     : {date}
Duration : {duration}
Speakers : {speakers}
"""


def _set_cell(table, row: int, col: int, text: str, bold: bool = False):
    cell = table.cell(row, col)
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def _add_section_heading(doc, text: str):
    from docx.shared import Pt, RGBColor
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.size = Pt(16)


def _add_markdown_content(doc, text: str):
    """Very basic markdown→docx rendering for the summary."""
    from docx.shared import Pt, RGBColor
    for line in text.split("\n"):
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)
